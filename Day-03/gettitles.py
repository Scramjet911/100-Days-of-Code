from bs4 import BeautifulSoup
import docx
from docx import Document
import requests
import sys

def remove_hyperlink(para):
    p = para._p
    label_r = p.xpath("./w:hyperlink/w:r")[0]
    hyperlink = p.xpath("./w:hyperlink")[0]
    if('list' not in str(type(label_r))):
        hyperlink.addprevious(label_r)
        p.remove(hyperlink)

def gettext(fileloc):
    finalText = []
    try:
        doc = Document(fileloc)
    except:
        return -1

    for line in doc.paragraphs:
        try:
            remove_hyperlink(line)
        except IndexError:
            pass
        if(line.text !=''):
            finalText.append(line.text)
    return finalText


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external = True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn("r:id"),r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rpr = docx.oxml.shared.OxmlElement('w:rpr')

    new_run.append(rpr)
    new_run.text = text
    hyperlink.append(new_run)

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000FF")
    rpr.append(c)

    para._p.append(hyperlink)

    return hyperlink

def get_h1(url):
    response = requests.get(url)
    content = BeautifulSoup(response.content,"html.parser")
    headText = content.find('h1')
    if(headText is not None):
        headText = headText.text
        return headText

argsList = sys.argv
inputfile = 'Input.docx'
outputfile = 'Output.docx'
if(len(argsList) > 3):
    print("Incorrect Number of arguments (Maximum 2 Arguments are allowed")

elif(len(argsList)==2):
    inputfile = sys.argv[1]
elif(len(argsList)==3):
    inputfile = sys.argv[1]
    outputfile = sys.argv[2]
    
docText = gettext(inputfile)
if(docText == -1):
    print("Input Document Not Found")
    exit()
newdoc = Document()
newdoc.add_heading('The Formatted Links List', level = 0)
for i in docText:
    h1 = get_h1(i)
    if(h1 == None):
        h1 = "Heading Not found"
    para = newdoc.add_paragraph(h1 + ' : ')
    add_hyperlink(para,'Open Site',i)

newdoc.save(outputfile)
print("Formatted Document saved as {}".format(outputfile))